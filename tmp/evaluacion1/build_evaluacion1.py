from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\cajgo\Downloads\1.5_GuiaEstudiante_Fase 1_Definicion Proyecto APT.docx")
OUTPUT = Path(r"C:\Users\cajgo\Documents\Codex\LocalitoSaaS\entregables\Evaluacion1\Definicion_Proyecto_APT_Localito_Camilo_Gonzalez.docx")

DARK_BLUE = "1F3864"
MID_BLUE = "4472C4"
LIGHT_BLUE = "D9E2F3"
GUIDE_BLUE = "548DD4"
BLACK = "000000"
WHITE = "FFFFFF"
GRID = "B7B7B7"
PHASE_1 = "E2F0D9"
PHASE_2 = "FFF2CC"
PHASE_3 = "FCE4D6"
ACTIVE = "4472C4"
PENDING = "FFF2CC"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, size=10, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_cell_text(cell, text: str, *, size=9.5, color=BLACK, bold=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, highlight=None):
    for p in cell.paragraphs[1:]:
        cell._element.remove(p._element)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    if highlight:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), highlight)
        run._element.get_or_add_rPr().append(shd)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
    return p


def set_cell_paragraphs(cell, paragraphs, *, size=9.5, color=BLACK):
    for p in list(cell.paragraphs)[1:]:
        cell._element.remove(p._element)
    p0 = cell.paragraphs[0]
    clear_paragraph(p0)
    all_paragraphs = [p0]
    for _ in range(max(0, len(paragraphs) - 1)):
        all_paragraphs.append(cell.add_paragraph())
    for p, item in zip(all_paragraphs, paragraphs):
        if isinstance(item, tuple):
            kind, text = item
        else:
            kind, text = "body", item
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if kind != "label" else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        if kind == "bullet":
            p.style = "List Paragraph"
            pPr = p._p.get_or_add_pPr()
            numPr = pPr.find(qn("w:numPr"))
            if numPr is None:
                numPr = OxmlElement("w:numPr")
                pPr.append(numPr)
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "3")
            numPr.append(ilvl)
            numPr.append(numId)
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=(kind == "label"))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=90, start=110, bottom=90, end=110)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def keep_table_rows(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def remove_row(table, row):
    table._tbl.remove(row._tr)


def insert_page_break_before_table(table):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    table._tbl.addprevious(p)


def style_body_table(table, header_rows=1, body_size=8.5):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=70, start=80, bottom=70, end=80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_font(r, size=body_size, color=BLACK, bold=(ri < header_rows))
        if ri < header_rows:
            for cell in row.cells:
                shade_cell(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if table.rows:
        repeat_header(table.rows[0])
    keep_table_rows(table)


def add_section_heading(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_text(cell, title, size=12.5, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    set_font(r, size=9.3, color=BLACK, italic=italic)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=8.7, color=BLACK)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    # Remove an invisible stray character present in the supplied guide.
    if doc.paragraphs and doc.paragraphs[0].text.strip() == "}":
        clear_paragraph(doc.paragraphs[0])
    # The supplied guide forces Part II onto a new page. Removing only the
    # earlier break lets the expanded Part I flow naturally instead of leaving
    # a nearly empty page when the competencies cell continues.
    first_page_break = doc.paragraphs[12]._p.find(".//" + qn("w:br"))
    if first_page_break is not None:
        first_page_break.getparent().remove(first_page_break)

    # 1. Antecedentes personales
    t = doc.tables[1]
    values = [
        "Camilo González",
        "[RUT PENDIENTE DE CONFIRMACIÓN]",
        "Ingeniería en Informática",
        "[SEDE PENDIENTE DE CONFIRMACIÓN]",
    ]
    for i, value in enumerate(values):
        set_cell_text(t.cell(i, 1), value, size=10, color=DARK_BLUE,
                      bold=(i == 0), align=WD_ALIGN_PARAGRAPH.LEFT,
                      highlight=PENDING if "PENDIENTE" in value else None)

    # 2. Descripción del Proyecto APT
    t = doc.tables[3]
    set_cell_text(
        t.cell(0, 1),
        "Localito: PWA SaaS multi-negocio para la gestión de comercios de barrio.",
        size=10, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_paragraphs(t.cell(1, 1), [
        ("bullet", "Desarrollo e implantación de soluciones de software."),
        ("bullet", "Diseño e implementación de modelos de datos."),
        ("bullet", "Gestión de proyectos informáticos."),
        ("bullet", "Aseguramiento de calidad y pruebas de software."),
    ], size=9.5, color=BLACK)
    set_cell_paragraphs(t.cell(2, 1), [
        ("bullet", "Realizar pruebas de certificación de productos y procesos."),
        ("bullet", "Gestionar proyectos informáticos y apoyar la toma de decisiones."),
        ("bullet", "Construir modelos de datos escalables para la organización."),
        ("bullet", "Desarrollar, integrar e implantar soluciones de software."),
    ], size=9.2, color=BLACK)

    # 3. Fundamentación
    t = doc.tables[5]
    set_cell_paragraphs(t.cell(0, 1), [
        "Localito nace de un dolor operacional observable en almacenes, botillerías, minimarkets y otros comercios de barrio chilenos: una parte importante de sus ventas, inventario, caja y cuentas fiadas se controla mediante cuadernos, memoria, planillas separadas o sistemas genéricos que no siempre se adaptan a una operación móvil y de baja complejidad. Esta fragmentación dificulta conocer el stock real, detectar diferencias de caja, recordar deudas, reponer oportunamente y tomar decisiones con información confiable.",
        "El cliente ideal inicial es un micro o pequeño comercio administrado directamente por su dueño, con operación manual o fragmentada y baja infraestructura tecnológica. La problemática afecta al dueño y a sus vendedores, porque aumenta el tiempo destinado a conciliaciones y búsquedas, y también puede afectar al cliente final mediante errores de precio, falta de productos o registros poco claros del fiado.",
        "La propuesta de valor es reunir ventas, inventario, clientes, fiado, caja, compras y proveedores en una PWA instalable y accesible desde el celular. Como diferenciadores, incorpora aislamiento de datos por negocio, funcionamiento controlado ante conectividad inestable, lectura de códigos de barras y reconocimiento visual opcional con confirmación humana. El modelo comercial es una suscripción mensual simple por negocio, todavía sujeta a validación mediante entrevistas, demostraciones y disposición a pagar. No se atribuyen ventas, ahorros ni retorno de inversión mientras no exista evidencia medida.",
        "Para la informática, el proyecto es relevante porque exige resolver seguridad, arquitectura multi-tenant, consistencia transaccional, experiencia móvil, pruebas, despliegue y trazabilidad; son desafíos propios del desarrollo profesional de software y de productos digitales vendibles.",
    ], size=9.2)
    set_cell_paragraphs(t.cell(1, 1), [
        "El proyecto busca desarrollar y validar técnica y comercialmente un producto mínimo viable de Localito. La solución permitirá administrar negocios separados entre sí; registrar productos, stock, ventas, clientes y fiados; controlar caja, compras y proveedores; y operar desde un navegador móvil mediante una PWA. El núcleo se complementará con roles de acceso, auditoría, cola local ante fallas de red, lectura de códigos y visión asistida opcional.",
        "La problemática se abordará mediante levantamiento de necesidades, definición de requisitos y criterios de aceptación, diseño de arquitectura y datos, desarrollo incremental, pruebas automatizadas y manuales, y validación con potenciales usuarios. La emisión tributaria ante el SII, múltiples sucursales complejas, facturación automática de la suscripción y Transbank productivo quedan fuera del MVP para proteger la factibilidad del semestre.",
    ], size=9.3)
    set_cell_paragraphs(t.cell(2, 1), [
        "Localito se relaciona directamente con el perfil de egreso de Ingeniería en Informática porque integra el ciclo completo de una solución: análisis de una necesidad, planificación, diseño de arquitectura, construcción de un modelo de datos, desarrollo full stack, integración de componentes, pruebas y preparación de despliegue.",
        "La competencia de pruebas se aplica al definir criterios de aceptación, automatizar reglas críticas y documentar casos manuales. La gestión de proyectos se evidencia en el alcance, backlog, sprints, riesgos, responsables y control de avances. El modelamiento de datos es necesario para separar negocios y mantener consistencia en ventas, stock, caja y fiados. El desarrollo e implantación de software se materializa en la PWA React, la API Node.js, PostgreSQL, seguridad por roles y configuración de despliegue. En conjunto, estas competencias permiten responder de forma trazable y mantenible a la problemática seleccionada.",
    ], size=9.3)
    set_cell_paragraphs(t.cell(3, 1), [
        "El proyecto se relaciona con mis intereses profesionales porque me permite participar en el diseño de una solución digital completa para un problema real. Me interesa profundizar en arquitectura de software, desarrollo backend, bases de datos, seguridad y construcción de productos SaaS escalables.",
        "Mi aporte principal se orienta a coordinar la arquitectura técnica, integrar la API con el modelo de datos, resguardar el aislamiento por negocio y mantener coherencia entre las reglas operacionales y los objetivos comerciales. Desarrollar Localito contribuirá a mi formación al obligarme a convertir necesidades de clientes en decisiones técnicas verificables, equilibrando alcance, calidad, seguridad y viabilidad comercial.",
    ], size=9.3)
    set_cell_paragraphs(t.cell(4, 1), [
        "El proyecto es factible durante un semestre de 18 semanas porque se trabajará sobre un MVP priorizado y un prototipo base, usando tecnologías disponibles para el equipo: computadores personales, teléfonos móviles, Git, Node.js, React, TypeScript, PostgreSQL y servicios de despliegue con planes académicos o gratuitos. El trabajo se distribuirá en iteraciones cortas con responsables definidos y evidencia verificable.",
        "Facilitadores: conocimiento previo de las tecnologías, repositorio existente, separación modular, datos de demostración y posibilidad de probar los flujos principales sin integraciones productivas. Dificultades posibles: crecimiento del alcance, acceso limitado a comerciantes, conectividad móvil, credenciales de servicios externos y tiempo de documentación. Se mitigarán congelando el alcance del MVP, usando alternativas controladas —código de barras o pista manual cuando no exista IA, Webpay simulado y modo local—, reservando semanas para integración y pruebas, y registrando resultados negativos o pendientes sin sobreestimarlos.",
    ], size=9.3)

    # 4. Objetivos
    t = doc.tables[7]
    set_cell_text(
        t.cell(0, 1),
        "Desarrollar y validar técnica y comercialmente una PWA SaaS multi-negocio que permita a pequeños comercios de barrio controlar ventas, inventario, clientes, fiados, caja y compras desde dispositivos móviles, aplicando seguridad, trazabilidad, pruebas y un modelo de suscripción sujeto a validación.",
        size=9.6,
    )
    set_cell_paragraphs(t.cell(1, 1), [
        ("bullet", "Analizar y validar las necesidades operacionales y comerciales del segmento objetivo mediante entrevistas y observación de procesos."),
        ("bullet", "Definir el alcance, los requisitos, los criterios de aceptación y el backlog priorizado del MVP."),
        ("bullet", "Diseñar e implementar una arquitectura multi-negocio y un modelo de datos escalable que resguarden la separación e integridad de la información."),
        ("bullet", "Implementar e integrar los módulos prioritarios, la experiencia PWA, la cola offline, los códigos de barras y la visión opcional con confirmación humana."),
        ("bullet", "Verificar la calidad, seguridad y usabilidad mediante pruebas automatizadas, casos manuales y evidencia reproducible."),
        ("bullet", "Evaluar la propuesta de valor, activación, uso y disposición a pagar, y documentar resultados, limitaciones y decisiones de mejora."),
    ], size=9.1)

    # 5. Metodología
    t = doc.tables[9]
    set_cell_paragraphs(t.cell(1, 0), [
        "Se utilizará una metodología ágil e incremental basada en Scrum, adaptada al contexto académico. El trabajo se organizará en un backlog priorizado por valor y riesgo, iteraciones cortas, revisión periódica de avances y retrospectivas. Cada historia o actividad tendrá responsable, criterio de aceptación y evidencia. Este enfoque es coherente con la necesidad de aprender de usuarios y ajustar el producto sin comprometer el alcance del MVP (Schwaber & Sutherland, 2020).",
        ("label", "Etapas de trabajo"),
        ("bullet", "Descubrimiento: entrevistas, observación del proceso actual, segmentación del cliente y validación del dolor."),
        ("bullet", "Definición: alcance, requisitos, backlog, arquitectura, modelo de datos, riesgos y métricas."),
        ("bullet", "Construcción incremental: desarrollo de módulos, integración continua, revisión de código y demostración al cierre de cada iteración."),
        ("bullet", "Verificación: typecheck, compilación, pruebas automatizadas, casos manuales, seguridad por roles, compatibilidad móvil y trazabilidad de defectos."),
        ("bullet", "Validación: demostraciones con potenciales usuarios y, si existe acceso, un piloto controlado; se medirán comprensión, activación, recurrencia, tiempo por venta, exactitud de inventario y disposición a pagar."),
        ("bullet", "Cierre: consolidación de evidencias, análisis de resultados, limitaciones, mejoras, informe y presentación."),
        ("label", "Funciones y responsabilidades"),
        ("bullet", "Camilo González: coordinación técnica, arquitectura, backend, API, modelo de datos, seguridad multi-negocio e integración."),
        ("bullet", "Alexander Patiño: frontend React, experiencia PWA, flujos móviles, accesibilidad e integración de interfaz con la API."),
        ("bullet", "Samuel Solis: planificación y seguimiento, matriz de pruebas, documentación, control de evidencias y apoyo a validación con usuarios."),
        "El equipo realizará revisiones cruzadas; la asignación define responsables principales, pero no elimina la colaboración ni la responsabilidad conjunta por la calidad.",
    ], size=9.1)

    # 6. Evidencias propuestas
    t = doc.tables[11]
    while len(t.rows) > 1:
        remove_row(t, t.rows[-1])
    evidence_rows = [
        ("Avance — Fase 1", "Definición y validación del problema", "Ficha de cliente ideal, entrevistas o registros de observación, mapa del proceso actual, dolor priorizado y síntesis de objeciones.", "Demuestra que la solución responde a un problema concreto y que la propuesta comercial se basa en evidencia, no solo en supuestos."),
        ("Avance — Fase 1", "Backlog, alcance y arquitectura", "Requisitos priorizados, criterios de aceptación, riesgos, diagrama de arquitectura y modelo de datos multi-negocio.", "Vincula las competencias de planificación, modelamiento de datos y diseño de software con decisiones trazables."),
        ("Avance — Fase 2", "Incremento funcional demostrable", "Repositorio, historial de versiones y demostración de autenticación, ventas, stock, fiado, caja, compras y roles.", "Permite observar construcción e integración real de componentes, sin confundir código existente con funcionalidad validada."),
        ("Avance — Fase 2", "Informe de calidad", "Salida de typecheck, compilación y pruebas; matriz de casos con ambiente, pasos, responsable, resultado y evidencia.", "Comprueba calidad disciplinaria y diferencia explícitamente entre casos aprobados, fallidos y pendientes."),
        ("Final — Fase 3", "Validación de usabilidad y propuesta de valor", "Resultados de demostraciones o piloto, tiempos observados, errores, comprensión, recurrencia, objeciones y disposición a pagar.", "Permite decidir si Localito reduce el dolor priorizado y si existe una base razonable para un producto vendible."),
        ("Final — Fase 3", "MVP desplegado y guía de reproducción", "URL o entorno de demostración, README, configuración sin secretos, datos demo y pasos para ejecutar los flujos principales.", "Acredita implantación reproducible y facilita que el docente verifique el producto."),
        ("Final — Fase 3", "Informe y presentación final", "Documento técnico-comercial, conclusiones, limitaciones, backlog futuro, presentación y demostración guiada.", "Integra el ciclo completo, comunica aprendizajes y sustenta decisiones con evidencias verificables."),
    ]
    for values in evidence_rows:
        cells = t.add_row().cells
        for ci, value in enumerate(values):
            set_cell_text(cells[ci], value, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT)
    note_cells = t.add_row().cells
    note_cell = note_cells[0]
    for c in note_cells[1:]:
        note_cell = note_cell.merge(c)
    set_cell_text(
        note_cell,
        "Nota: estas evidencias constituyen una propuesta inicial y deben validarse con el docente antes de fijar su versión definitiva.",
        size=8.2, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    shade_cell(note_cell, PENDING)
    style_body_table(t, header_rows=1, body_size=8.2)

    # 7. Plan de trabajo
    t = doc.tables[13]
    while len(t.rows) > 2:
        remove_row(t, t.rows[-1])
    plan_rows = [
        ("Gestión de proyectos / pruebas", "Descubrimiento y validación", "Entrevistar u observar comercios, definir cliente ideal, dolor, alternativas y criterios de éxito.", "Guion, consentimiento, notas, planilla de síntesis.", "Semanas 1–3", "Samuel / equipo", "Acceso a comercios puede ser limitado; usar contactos, entrevistas remotas y registrar rechazos."),
        ("Gestión de proyectos", "Alcance y backlog", "Priorizar el MVP, requisitos, historias, criterios de aceptación, riesgos y métricas.", "Backlog, Git, tablero y documentación.", "Semanas 2–4", "Samuel y Camilo", "Evitar crecimiento de alcance; congelar núcleo y mantener roadmap separado."),
        ("Modelos de datos / software", "Arquitectura y datos", "Diseñar aislamiento por negocio, entidades, relaciones, permisos y contratos de API.", "PostgreSQL, diagramas, TypeScript y revisión técnica.", "Semanas 3–6", "Camilo", "Riesgo de inconsistencias; usar restricciones, transacciones y revisión cruzada."),
        ("Desarrollo de software", "Núcleo operacional", "Consolidar autenticación, productos, inventario, ventas, clientes, fiado, caja, compras y proveedores.", "Node.js, React, PostgreSQL, entorno local.", "Semanas 5–10", "Camilo y Alexander", "Dependencias entre módulos; integrar por incrementos y demostrar cada flujo."),
        ("Integración e implantación", "Experiencia PWA e integraciones", "Optimizar uso móvil, cola offline, códigos, visión opcional, correo y pago demostrativo.", "PWA, móvil, ZXing, servicios externos y entorno HTTPS.", "Semanas 8–12", "Alexander y Camilo", "Credenciales o conectividad pueden fallar; conservar flujos alternativos controlados."),
        ("Pruebas de certificación", "Verificación y mejora", "Ejecutar pruebas automatizadas y manuales, registrar defectos, corregir y repetir casos críticos.", "Matriz, TypeScript, runner de pruebas, navegadores y móviles.", "Semanas 8–15", "Samuel / equipo", "No elevar casos a aprobados sin salida o captura verificable."),
        ("Gestión / pruebas", "Validación comercial", "Probar mensaje, demostración, activación y disposición a pagar; realizar piloto acotado si existe acceso.", "Guion, prototipo, planilla de métricas y consentimiento.", "Semanas 12–16", "Samuel y Alexander", "La muestra puede ser pequeña; reportar alcance, sesgos y resultados negativos."),
        ("Implantación / gestión", "Cierre y defensa", "Estabilizar demo, consolidar evidencias, informe, conclusiones, limitaciones y presentación.", "Repositorio, hosting, documentos, capturas y presentación.", "Semanas 16–18", "Equipo completo", "Reservar contingencia y usar un guion reproducible con datos de demostración."),
    ]
    for values in plan_rows:
        cells = t.add_row().cells
        for ci, value in enumerate(values):
            set_cell_text(cells[ci], value, size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    style_body_table(t, header_rows=2, body_size=7.5)
    repeat_header(t.rows[1])

    # Keep each major planning block with its table instead of orphaning a
    # section heading at the bottom of the preceding page.
    insert_page_break_before_table(doc.tables[12])
    insert_page_break_before_table(doc.tables[14])

    # 8. Carta Gantt: correct the duplicate S16 and use exactly 18 weeks.
    old_gantt = doc.tables[15]
    old_gantt._element.getparent().remove(old_gantt._element)
    gantt = doc.add_table(rows=10, cols=19)
    gantt.style = "Table Grid"
    gantt.alignment = WD_TABLE_ALIGNMENT.CENTER
    gantt.autofit = False
    gantt.cell(0, 0).merge(gantt.cell(1, 0))
    gantt.cell(0, 1).merge(gantt.cell(0, 4))
    gantt.cell(0, 5).merge(gantt.cell(0, 15))
    gantt.cell(0, 16).merge(gantt.cell(0, 18))
    set_cell_text(gantt.cell(0, 0), "Actividad", size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for start, title, fill in ((1, "Fase 1", PHASE_1), (5, "Fase 2", PHASE_2), (16, "Fase 3", PHASE_3)):
        set_cell_text(gantt.cell(0, start), title, size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(gantt.cell(0, start), fill)
    for week in range(1, 19):
        cell = gantt.cell(1, week)
        set_cell_text(cell, f"S{week}", size=6.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, PHASE_1 if week <= 4 else PHASE_2 if week <= 15 else PHASE_3)
    activities = [
        ("Problema y usuarios", 1, 3),
        ("Alcance y backlog", 2, 4),
        ("Arquitectura y datos", 3, 6),
        ("Núcleo operacional", 5, 10),
        ("PWA e integraciones", 8, 12),
        ("Pruebas y mejoras", 8, 15),
        ("Validación comercial", 12, 16),
        ("Cierre y defensa", 16, 18),
    ]
    for ri, (activity, start, end) in enumerate(activities, start=2):
        set_cell_text(gantt.cell(ri, 0), activity, size=7.2, align=WD_ALIGN_PARAGRAPH.LEFT)
        for week in range(1, 19):
            cell = gantt.cell(ri, week)
            set_cell_text(cell, "X" if start <= week <= end else "", size=6.8,
                          color=WHITE if start <= week <= end else BLACK,
                          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            if start <= week <= end:
                shade_cell(cell, ACTIVE)
    for row in gantt.rows:
        row.cells[0].width = Inches(1.4)
        for cell in row.cells[1:]:
            cell.width = Inches(0.25)
        for cell in row.cells:
            set_cell_border(cell, color="A6A6A6", size=4)
            set_cell_margins(cell, top=40, start=25, bottom=40, end=25)
    repeat_header(gantt.rows[0])
    repeat_header(gantt.rows[1])
    keep_table_rows(gantt)

    # Complementary rubric content retained after the official guide.
    doc.add_page_break()
    add_section_heading(doc, "9. Resumen y Abstract")
    p = doc.add_paragraph()
    r = p.add_run("Resumen")
    set_font(r, size=11, color=DARK_BLUE, bold=True)
    add_body(doc, "Localito es una PWA SaaS multi-negocio orientada a pequeños comercios de barrio que necesitan mejorar el control de ventas, inventario, clientes, fiados, caja, compras y proveedores. La propuesta prioriza una experiencia móvil de baja fricción, aislamiento seguro de datos, funcionamiento controlado ante conectividad inestable, códigos de barras y reconocimiento visual opcional. El proyecto se desarrollará con metodología ágil, pruebas y evidencia trazable. Su modelo de suscripción, beneficios y disposición a pagar serán tratados como hipótesis hasta validarlos con entrevistas, demostraciones y un piloto acotado.")
    p = doc.add_paragraph()
    r = p.add_run("Abstract")
    set_font(r, size=11, color=DARK_BLUE, bold=True)
    add_body(doc, "Localito is a multi-business SaaS Progressive Web Application for small neighborhood stores that need better control of sales, inventory, customers, store credit, cash, purchases, and suppliers. The proposal emphasizes a low-friction mobile experience, secure data isolation, controlled operation during unstable connectivity, barcode scanning, and optional visual recognition. The project will follow an agile methodology supported by testing and traceable evidence. Its subscription model, expected benefits, and willingness to pay will remain hypotheses until they are validated through interviews, demonstrations, and a limited pilot.")

    add_section_heading(doc, "10. Conclusions")
    add_body(doc, "Localito is a relevant and feasible capstone proposal because it addresses a concrete operational problem and requires the integration of project management, data modeling, software development, testing, and deployment competencies. Its main strength is not the number of features, but the connection between a specific customer pain and a mobile solution that can be demonstrated and measured. The project must preserve a controlled MVP and avoid claiming sales, savings, or customer impact without evidence. If the team validates usability, recurring use, and willingness to pay, Localito could evolve from an academic prototype into a sustainable software product.")

    add_section_heading(doc, "11. Reflection")
    add_body(doc, "Defining Localito has helped me understand that a technically complete system is not automatically a valuable product. As a future software engineer, I need to connect architecture, security, and data decisions with the real needs of users. My main challenge will be balancing technical ambition with the time available and collecting honest evidence, including negative results. This project will allow me to strengthen backend development, multi-business architecture, and technical coordination while learning to validate commercial assumptions before treating them as facts.")

    add_section_heading(doc, "12. Referencias")
    add_reference(doc, "Localito. (2026). Documento de Proyecto — Localito (versión 1.0) [Documento interno del proyecto].")
    add_reference(doc, "Localito. (2026). Matriz de pruebas funcionales — Localito [Documento interno del proyecto].")
    add_reference(doc, "Project Management Institute. (2021). A guide to the project management body of knowledge (PMBOK® Guide) (7th ed.).")
    add_reference(doc, "Schwaber, K., & Sutherland, J. (2020). The Scrum Guide: The definitive guide to Scrum: The rules of the game.")
    add_reference(doc, "Sommerville, I. (2016). Software engineering (10th ed.). Pearson.")

    # Global cleanup: no fixed heights; keep institutional styling and use readable body text.
    for table in doc.tables:
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            for h in list(trPr.findall(qn("w:trHeight"))):
                trPr.remove(h)

    core = doc.core_properties
    core.title = "Definición Proyecto APT — Localito"
    core.subject = "Evaluación 1 — Asignatura Capstone"
    core.author = "Camilo González"
    core.keywords = "Localito, Capstone, PWA, SaaS, comercio de barrio"
    core.comments = "Documento completado a partir de la guía institucional, sin modificar el archivo original."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
