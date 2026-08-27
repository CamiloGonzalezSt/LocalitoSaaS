from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import sys

OUT=Path(sys.argv[1]); OUT.parent.mkdir(parents=True,exist_ok=True)
d=Document(); s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1); s.header_distance=s.footer_distance=Inches(.49); s.different_first_page_header_footer=True
BLUE=RGBColor(31,77,120); NAVY=RGBColor(32,55,72); GOLD=RGBColor(154,112,30); GRAY=RGBColor(90,90,90)

def rf(r,size=11,bold=None,italic=None,color=None):
 r.font.name='Arial'; r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'),'Arial'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); r.font.size=Pt(size)
 if bold is not None:r.bold=bold
 if italic is not None:r.italic=italic
 if color:r.font.color.rgb=color
def p(text='',style=None,after=8,before=0,align=WD_ALIGN_PARAGRAPH.JUSTIFY,bold=False,italic=False,size=11,color=None):
 x=d.add_paragraph(style=style); x.alignment=align; x.paragraph_format.space_before=Pt(before); x.paragraph_format.space_after=Pt(after); x.paragraph_format.line_spacing=1.5
 if text:rf(x.add_run(text),size,bold,italic,color)
 return x
def h(text,level=1): return p(text,f'Heading {level}',after=10 if level==1 else 6,before=18 if level==1 else 12,align=WD_ALIGN_PARAGRAPH.LEFT)
def bullet(text):
 x=d.add_paragraph(style='List Bullet'); x.paragraph_format.left_indent=Inches(.375); x.paragraph_format.first_line_indent=Inches(-.194); x.paragraph_format.space_after=Pt(4); x.paragraph_format.line_spacing=1.5; rf(x.add_run(text)); return x
def shade(cell,fill='F4F6F9'):
 sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(sh)
def set_cell_margins(cell):
 tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
 if m is None:m=OxmlElement('w:tcMar');tc.append(m)
 for side,val in [('top',80),('bottom',80),('start',120),('end',120)]:
  e=m.find(qn('w:'+side));
  if e is None:e=OxmlElement('w:'+side);m.append(e)
  e.set(qn('w:w'),str(val));e.set(qn('w:type'),'dxa')
def tbl(headers,rows,widths):
 t=d.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.autofit=False
 total=9360; grid=t._tbl.tblGrid
 for e in list(grid):grid.remove(e)
 for w in widths:
  gc=OxmlElement('w:gridCol');gc.set(qn('w:w'),str(w));grid.append(gc)
 pr=t._tbl.tblPr; tw=pr.first_child_found_in('w:tblW');tw.set(qn('w:type'),'dxa');tw.set(qn('w:w'),str(total)); ind=OxmlElement('w:tblInd');ind.set(qn('w:w'),'120');ind.set(qn('w:type'),'dxa');pr.append(ind)
 for i,v in enumerate(headers):t.cell(0,i).text=v;shade(t.cell(0,i));set_cell_margins(t.cell(0,i))
 trPr=t.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader');hdr.set(qn('w:val'),'true');trPr.append(hdr)
 for row in rows:
  cs=t.add_row().cells
  for i,v in enumerate(row):cs[i].text=str(v);set_cell_margins(cs[i])
 for row in t.rows:
  for i,c in enumerate(row.cells):
   tcw=c._tc.get_or_add_tcPr().first_child_found_in('w:tcW');tcw.set(qn('w:w'),str(widths[i]));tcw.set(qn('w:type'),'dxa')
   for pp in c.paragraphs:
    pp.paragraph_format.space_after=Pt(2);pp.paragraph_format.line_spacing=1.15
    for r in pp.runs:rf(r,9,bold=(row is t.rows[0]))
 p('',after=2);return t
def field(par,code):
 r=par.add_run(); b=OxmlElement('w:fldChar');b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText');i.set(qn('xml:space'),'preserve');i.text=code; e=OxmlElement('w:fldChar');e.set(qn('w:fldCharType'),'end');r._r.extend([b,i,e]);rf(r,9,color=GRAY)

# styles
normal=d.styles['Normal'];normal.font.name='Arial';normal.font.size=Pt(11);normal._element.rPr.rFonts.set(qn('w:ascii'),'Arial');normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial');normal.paragraph_format.line_spacing=1.5;normal.paragraph_format.space_after=Pt(8)
for name,size,col,bef,aft in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,12,6),('Heading 3',12,NAVY,8,4)]:
 st=d.styles[name];st.font.name='Arial';st.font.size=Pt(size);st.font.bold=True;st.font.color.rgb=col;st._element.rPr.rFonts.set(qn('w:ascii'),'Arial');st._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial');st.paragraph_format.space_before=Pt(bef);st.paragraph_format.space_after=Pt(aft);st.paragraph_format.keep_with_next=True

# running furniture
hdr=s.header.paragraphs[0];hdr.alignment=WD_ALIGN_PARAGRAPH.LEFT;rf(hdr.add_run('LOCALITO  |  INFORME DE AUTOEVALUACIÓN APT'),9,True,color=GRAY)
f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;rf(f.add_run('Capstone · Fase 1  |  Página '),9,color=GRAY);field(f,'PAGE')

# editorial cover
p('',after=110)
p('PROYECTO APT · FASE 1',after=18,align=WD_ALIGN_PARAGRAPH.CENTER,bold=True,size=10,color=GOLD)
p('Informe de Autoevaluación',after=8,align=WD_ALIGN_PARAGRAPH.CENTER,bold=True,size=28,color=NAVY)
p('Localito',after=4,align=WD_ALIGN_PARAGRAPH.CENTER,bold=True,size=22,color=BLUE)
p('PWA SaaS multi-negocio para comercios de barrio',after=58,align=WD_ALIGN_PARAGRAPH.CENTER,size=14,color=NAVY)
p('Estudiante: [NOMBRE COMPLETO]',after=4,align=WD_ALIGN_PARAGRAPH.CENTER,bold=True,size=11)
p('Carrera: [NOMBRE OFICIAL DE LA CARRERA]',after=4,align=WD_ALIGN_PARAGRAPH.CENTER,size=11)
p('Asignatura: Capstone (PTY4614)',after=4,align=WD_ALIGN_PARAGRAPH.CENTER,size=11)
p('Fecha: [FECHA DE ENTREGA]',after=4,align=WD_ALIGN_PARAGRAPH.CENTER,size=11)
d.add_page_break()

h('Índice')
for n,title in [('1','Resumen y Abstract'),('2','Descripción y relevancia del Proyecto APT'),('3','Relación con el perfil de egreso'),('4','Intereses profesionales'),('5','Factibilidad'),('6','Objetivos'),('7','Metodología'),('8','Plan de trabajo'),('9','Evidencias'),('10','Aspectos formales e indicadores de calidad'),('11','Autoevaluación razonada'),('12','Conclusions'),('13','Reflection'),('14','Referencias'),('Anexo A','Estado de validación')]:p(f'{n}. {title}',after=3,align=WD_ALIGN_PARAGRAPH.LEFT)
p('Nota: los números de página se actualizan automáticamente al abrir el archivo en Microsoft Word.',italic=True,size=9,color=GRAY,align=WD_ALIGN_PARAGRAPH.LEFT)
d.add_page_break()

h('1. Resumen y Abstract')
h('Resumen',2)
p('Localito es una aplicación web progresiva (PWA) de tipo SaaS y arquitectura multi-negocio, orientada a almacenes y comercios de barrio. Su propósito es digitalizar ventas, inventario, caja, compras, proveedores, clientes y cuentas fiadas desde un teléfono móvil. Incluye control de acceso por roles, aislamiento de datos por negocio, operación sin conexión para flujos críticos y reconocimiento de productos mediante código de barras o visión asistida por inteligencia artificial. La iteración actual implementa el núcleo operacional y excluye deliberadamente la emisión tributaria ante el SII y la activación productiva de Transbank. El proyecto integra análisis de requerimientos, modelamiento de datos, desarrollo full stack, seguridad y pruebas, por lo que resulta pertinente para el perfil profesional de informática.')
p('Palabras clave: PWA, SaaS, comercio de barrio, inventario, punto de venta, multi-tenant, inteligencia artificial.',italic=True,size=10)
h('Abstract',2)
p('Localito is a multi-business Software as a Service Progressive Web Application designed for neighborhood stores and small retailers. Its purpose is to digitize sales, inventory, cash shifts, purchases, suppliers, customers, and store credit from a mobile phone. The solution includes role-based access control, tenant data isolation, offline support for critical workflows, and product recognition through barcodes or AI-assisted computer vision. The current iteration implements the operational core, while Chilean tax-compliant receipts and the production activation of Transbank remain outside its scope. The project integrates requirements analysis, data modeling, full-stack development, security, and testing, making it relevant to the professional profile of an information technology graduate.')
p('Keywords: PWA, SaaS, neighborhood retail, inventory, point of sale, multi-tenant, artificial intelligence.',italic=True,size=10)

h('2. Descripción y relevancia del Proyecto APT')
p('Localito funciona como una “caja inteligente de bolsillo”. Permite administrar negocios separados entre sí y ofrece, según el rol, creación de locales y usuarios, catálogo de productos, stock y kardex, alertas, ventas con descuentos y pagos divididos, fiado, devoluciones, proveedores, órdenes de compra, caja por turno, auditoría e intercambio CSV. También incorpora una PWA instalable, cola local para operaciones sin conexión, lectura de códigos de barra y un flujo opcional de reconocimiento de envases mediante IA cuando existe una clave configurada.')
p('La relevancia laboral radica en que reproduce desafíos reales: levantamiento de necesidades, diseño escalable, separación segura de datos, integración de interfaz, API y base de datos, reglas transaccionales, pruebas y despliegue. Su beneficio esperado es mejorar la trazabilidad de pequeños comercios y reducir errores de stock, caja y cuentas por cobrar. El ticket del MVP es un comprobante interno y no reemplaza una boleta tributaria.')

h('3. Relación con el perfil de egreso')
tbl(['Competencia','Aplicación','Evidencia'],[
['Pruebas y certificación','Casos funcionales, automatización y registro de pendientes manuales.','Matriz CP-01 a CP-46 y pruebas de reglas críticas.'],['Gestión de proyectos','Priorización del MVP, alcance, riesgos y decisiones técnicas.','Plan de trabajo, backlog y documentación.'],['Modelamiento de datos','Modelo relacional multi-negocio para la operación.','Esquema PostgreSQL y tenant derivado de la sesión.'],['Desarrollo de software','Integración de React PWA, API Node.js y PostgreSQL.','Código, compilación, endpoints y flujos.'],['Implantación','Configuración, despliegue HTTPS y puesta en marcha.','README, scripts y configuración de Vercel/PostgreSQL.']],[2500,3700,3160])

h('4. Relación con mis intereses profesionales')
p('El proyecto se relaciona con mi interés por desarrollar soluciones digitales completas que resuelvan problemas cotidianos y tengan impacto visible. Me permite profundizar en desarrollo full stack, PWA móviles, diseño seguro de sistemas SaaS, automatización de pruebas e integración responsable de inteligencia artificial. También fortalece mi capacidad para traducir una necesidad de negocio en una arquitectura y un producto demostrable. Este apartado debe complementarse durante la presentación con experiencias personales auténticas del/de la estudiante, sin atribuir antecedentes no acreditados.')

h('5. Factibilidad')
p('Localito es factible dentro de la asignatura porque el núcleo se divide en módulos verificables y usa tecnologías disponibles: Node.js, React, PostgreSQL, navegador móvil y herramientas de desarrollo estándar. El alcance se controla mediante un MVP: SII, múltiples sucursales complejas, facturación de la suscripción y Transbank productivo quedan fuera de esta iteración. La IA visual es opcional; sin clave, el sistema conserva un flujo controlado mediante código de barras o pista manual.')
tbl(['Factor','Riesgo','Tratamiento'],[['Tiempo','Crecimiento del alcance.','Congelar MVP y priorizar defectos/evidencia.'],['Recursos','Dependencia de PostgreSQL, HTTPS y servicios.','Modo demo y configuración documentada.'],['Técnica','Integración, concurrencia y offline.','Regresión, idempotencia, logs y matriz manual.'],['Seguridad','Credenciales o permisos incorrectos.','Variables privadas y pruebas de autorización.'],['Validación','Casos manuales sin evidencia.','Ejecutar, capturar y registrar resultado real.']],[1900,3160,4300])

h('6. Objetivos')
h('Objetivo general',2)
p('Desarrollar y validar una PWA SaaS multi-negocio que apoye la gestión de ventas, inventario, caja, compras, clientes y fiados de pequeños comercios, incorporando reconocimiento de productos desde dispositivos móviles y controles de seguridad y trazabilidad.')
h('Objetivos específicos',2)
for x in ['Analizar y documentar necesidades y límites del MVP.','Diseñar e implementar un modelo de datos escalable y aislado por negocio.','Integrar autenticación, productos, stock, ventas, fiado, caja, compras y auditoría.','Incorporar códigos de barra y reconocimiento visual opcional con confirmación humana.','Validar reglas críticas con pruebas automatizadas y ejecutar pruebas manuales funcionales y móviles.','Preparar evidencia técnica y una demostración reproducible, diferenciando lo implementado, simulado y pendiente.']:bullet(x)

h('7. Metodología de trabajo')
p('Se utiliza una metodología ágil e incremental basada en Scrum, adaptada a un proyecto académico individual. El backlog se ordena por valor y riesgo; cada iteración selecciona historias y criterios de aceptación, implementa un incremento, ejecuta revisión técnica y actualiza la evidencia. La calidad se aborda con control de versiones, validación de tipos, compilación, pruebas de API, pruebas manuales y registro de defectos. Cada requisito mantiene trazabilidad con su módulo, caso de prueba y evidencia. Las integraciones externas se validan en modo de demostración o prueba y no se presentan como productivas sin credenciales ni confirmación real.')

h('8. Plan de trabajo')
tbl(['Etapa','Actividades','Duración','Resultado'],[['Definición','Problema, usuarios, alcance, objetivos y backlog.','1 semana','Definición y backlog.'],['Diseño','Arquitectura, datos, seguridad e interfaz.','1 semana','Diseño trazable.'],['Núcleo','Autenticación, inventario, ventas y fiado.','2 semanas','Flujos centrales.'],['Operación','Caja, compras, devoluciones, CSV y offline.','2 semanas','MVP operacional.'],['Integraciones','Código de barras, IA opcional y Webpay demo.','1 semana','Flujos demostrables.'],['Validación','Typecheck, build, pruebas y casos manuales.','1 semana','Resultados registrados.'],['Cierre','Correcciones, memoria, capturas y ensayo.','1 semana','Entrega reproducible.']],[1700,4200,1400,2060])
p('Recursos principales: computador personal, Node.js, React, PostgreSQL, Git, navegador de escritorio y dispositivo móvil. Facilitadores: arquitectura modular, documentación existente y pruebas automatizadas. Obstaculizadores: dependencias externas, evidencia móvil pendiente y riesgo de expansión del alcance.',size=10)

h('9. Evidencias propuestas')
tbl(['Evidencia','Qué demuestra','Criterio'],[['Repositorio y código','Construcción e integración.','Módulos claros y sin secretos.'],['Esquema PostgreSQL','Modelo y aislamiento.','Entidades y relaciones coherentes.'],['Typecheck, build y tests','Calidad técnica.','Comandos sin errores y salida fechada.'],['Matriz CP-01 a CP-46','Cobertura y trazabilidad.','Estado real por caso; sin aprobación ficticia.'],['Capturas o video','Operación visible.','Paso, rol y resultado observables.'],['Pruebas de permisos','Seguridad multi-tenant.','403 y tenant derivado de sesión.'],['Documentación','Alcance honesto.','SII y Transbank productivo fuera de alcance.']],[2600,3380,3380])
p('Las pruebas automatizadas aprobadas pueden presentarse como evidencia disponible. Los casos marcados “Pendiente evidencia” deben ejecutarse antes de declararlos aprobados.',italic=True,size=10)

h('10. Aspectos formales e indicadores de calidad')
h('Aspectos formales',2)
p('El informe emplea Arial 11, interlineado 1,5, tamaño carta, márgenes de una pulgada, jerarquía de títulos, encabezado, pie y referencias. El resumen y el abstract se presentan en español e inglés; las conclusiones y la reflexión están redactadas únicamente en inglés, según la pauta.')
tbl(['Indicador disciplinario','Presencia','Estado'],[['1.1 Diseña pruebas','Sí','Matriz y criterios definidos.'],['1.2 Aplica pruebas','Parcial','Automatizadas aprobadas; manuales pendientes.'],['1.3 Mejora según resultados','Parcial','Requiere bitácora consolidada.'],['2.1 Planifica proyectos','Sí','Alcance, etapas, riesgos y recursos.'],['2.2 Controla proyectos','Parcial','Debe conservar evidencia periódica.'],['3.1 Diseña modelos de datos','Sí','Modelo multi-negocio.'],['3.2 Implementa modelos','Sí','Esquema PostgreSQL.'],['4.1 Construye software','Sí','Núcleo operacional.'],['4.2 Integra componentes','Sí','PWA, API y datos.'],['4.3 Implanta software','Parcial','Despliegue preparado; depende de configuración.']],[4200,1200,3960])

h('11. Autoevaluación razonada')
tbl(['IE','Nivel','Justificación / mejora'],[['1','Completamente logrado','Descripción, impacto y relevancia profesional.'],['2','Completamente logrado','Competencias relacionadas con evidencias.'],['3','Logrado','Conexión profesional clara; falta personalización auténtica.'],['4','Completamente logrado','Tiempo, recursos, riesgos y mitigaciones.'],['5','Completamente logrado','Objetivos claros y coherentes.'],['6','Completamente logrado','Metodología pertinente y trazable.'],['7','Completamente logrado','Plan con actividades, duración y recursos.'],['8','Completamente logrado','Evidencias justificadas con criterios.'],['9','Completamente logrado','Redacción, estructura y referencias revisadas.'],['10','Completamente logrado','Formato institucional solicitado aplicado.'],['11','Logrado','Diseño cubierto; validación/implantación aún parcial.'],['12','Completamente logrado','Requested sections communicate in upper-intermediate English.']],[700,2200,6460])
p('La autoevaluación no busca maximizar artificialmente el puntaje, sino identificar mejoras concretas. Las prioridades son completar evidencias manuales, consolidar la bitácora de defectos y registrar el resultado del despliegue configurado.',italic=True,size=10)

h('12. Conclusions')
p('Localito is a feasible and professionally relevant capstone project because it connects a real operational problem with an integrated software solution. Its scope is controlled through a clear MVP and separates implemented features from simulated or future integrations. The project demonstrates competencies in data modeling, full-stack development, security, project planning, and software testing. Its strongest point is the operational breadth of the implemented core. The main remaining challenge is to complete the manual validation evidence, document defects and improvements, and rehearse a reproducible demonstration. These actions will strengthen the credibility of the final presentation and prevent unverified results from being reported as completed.')

h('13. Reflection')
p('Working on Localito shows me that software quality depends on more than writing code. A useful solution needs a well-defined problem, realistic scope, secure data handling, testable requirements, and honest evidence. I have learned that external services such as payment platforms and AI models must be treated as dependencies with costs, credentials, privacy concerns, and fallback strategies. I also understand the importance of separating an academic demonstration from a production-ready service. My next priority is to execute the pending manual test cases, record the real outcomes, correct any defects, and explain the project in clear language for both technical and non-technical audiences.')

h('14. Referencias')
p('Localito. (2026). Documento de Proyecto — Localito (versión 1.0) [Documento interno del proyecto].',align=WD_ALIGN_PARAGRAPH.LEFT)
p('Localito. (2026). Matriz de pruebas funcionales — Localito [Documento interno del proyecto].',align=WD_ALIGN_PARAGRAPH.LEFT)
p('Localito. (2026). README del repositorio Localito [Documentación técnica del proyecto].',align=WD_ALIGN_PARAGRAPH.LEFT)

h('Anexo A. Estado resumido de validación')
tbl(['Estado','Alcance'],[['Automatizada aprobada','Sesiones y contraseñas, idempotencia, crédito, devoluciones, caja, compras, costo promedio e inventario, según la matriz local.'],['Pendiente evidencia','Flujos manuales de interfaz, móvil, PWA, cámara, CSV y varios permisos que aún requieren captura o registro.'],['Fuera de alcance','SII/boleta electrónica, Transbank productivo, múltiples sucursales complejas y facturación de la suscripción SaaS.']],[2300,7060])
p('Este anexo debe actualizarse con fecha, responsable, ambiente y evidencia antes de la defensa. Un caso pendiente no debe presentarse como aprobado.',italic=True,size=10)

# update fields on open
settings=d.settings._element; upd=OxmlElement('w:updateFields');upd.set(qn('w:val'),'true');settings.append(upd)
d.core_properties.title='Informe de Autoevaluación del Proyecto APT Localito';d.core_properties.subject='Capstone Fase 1';d.core_properties.author='[NOMBRE COMPLETO]';d.core_properties.keywords='Localito, PWA, SaaS, APT, autoevaluación'
d.save(OUT)
