from docx import Document
import json, sys
d=Document(sys.argv[1]); z={"paragraphs":[p.text for p in d.paragraphs],"tables":[]}
for t in d.tables:
 z["tables"].append([[c.text for c in r.cells] for r in t.rows])
print(json.dumps(z,ensure_ascii=False,indent=2))
