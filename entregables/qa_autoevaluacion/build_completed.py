from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import shutil, sys

src=Path(sys.argv[1]); out=Path(sys.argv[2]); out.parent.mkdir(parents=True,exist_ok=True)
shutil.copy2(src,out); d=Document(out)

def font(run,size=10,bold=False,italic=False):
 run.font.name='Calibri'; run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'),'Calibri'); run.font.size=Pt(size); run.bold=bold; run.italic=italic
def para(text='',size=10,bold=False,italic=False,align=None,space=4):
 p=d.add_paragraph(); p.paragraph_format.space_after=Pt(space); p.paragraph_format.line_spacing=1.08
 if align is not None:p.alignment=align
 r=p.add_run(text); font(r,size,bold,italic); return p
def heading(text,level=1):
 p=d.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4); p.paragraph_format.keep_with_next=True
 r=p.add_run(text); font(r,14 if level==1 else 11,bold=True); return p
def bullet(text):
 p=d.add_paragraph(style='List Paragraph'); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.first_line_indent=Cm(-0.3); p.paragraph_format.space_after=Pt(2); r=p.add_run('• '+text); font(r,10); return p
def shade(cell,fill):
 tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def table(headers, rows, widths=None):
 t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.autofit=False
 for i,h in enumerate(headers):
  c=t.rows[0].cells[i]; c.text=h; shade(c,'D9EAF7')
  for r in c.paragraphs[0].runs: font(r,9,bold=True)
 for row in rows:
  cells=t.add_row().cells
  for i,val in enumerate(row):
   cells[i].text=str(val)
   for p in cells[i].paragraphs:
    p.paragraph_format.space_after=Pt(1)
    for r in p.runs:font(r,8.5)
 if widths:
  for row in t.rows:
   for c,w in zip(row.cells,widths):c.width=Cm(w)
 d.add_paragraph().paragraph_format.space_after=Pt(2); return t

p=para('INFORME DE DEFINICIÓN Y AUTOEVALUACIÓN DEL PROYECTO APT',16,True,align=WD_ALIGN_PARAGRAPH.CENTER,space=6)
para('Proyecto: Localito — PWA SaaS multi-negocio para comercios de barrio',12,True,align=WD_ALIGN_PARAGRAPH.CENTER,space=8)
table(['Dato','Información'],[['Estudiante','[NOMBRE COMPLETO DEL/DE LA ESTUDIANTE]'],['Carrera','[NOMBRE OFICIAL DE LA CARRERA]'],['Asignatura','Capstone (PTY4614)'],['Fase','Fase 1 — Definición del Proyecto APT'],['Fecha','[FECHA DE ENTREGA]']], [5,15])

heading('Resumen',1)
para('Localito es una aplicación web progresiva (PWA) de tipo SaaS y arquitectura multi-negocio, orientada a almacenes y comercios de barrio. Su propósito es digitalizar ventas, inventario, caja, compras, proveedores, clientes y cuentas fiadas desde un teléfono móvil. La solución incorpora control de acceso por roles, aislamiento de datos por negocio, operación sin conexión para flujos críticos y reconocimiento de productos mediante código de barras o visión asistida por inteligencia artificial. El proyecto responde a procesos que todavía se registran en cuadernos, memoria o planillas dispersas, y propone una herramienta de baja fricción, instalable desde el navegador. La iteración actual implementa el núcleo operacional y excluye deliberadamente la emisión tributaria ante el SII y la activación productiva de Transbank. Su desarrollo integra análisis de requerimientos, diseño de datos, construcción de software, seguridad y pruebas, por lo que es pertinente para el perfil profesional de informática.')
para('Palabras clave: PWA, SaaS, comercio de barrio, inventario, punto de venta, multi-tenant, inteligencia artificial.',9,italic=True)

heading('Abstract',1)
para('Localito is a multi-business Software as a Service Progressive Web Application designed for neighborhood stores and small retailers. Its purpose is to digitize sales, inventory, cash shifts, purchases, suppliers, customers, and store credit from a mobile phone. The solution includes role-based access control, tenant data isolation, offline support for critical workflows, and product recognition through barcodes or AI-assisted computer vision. It addresses processes that are still managed with notebooks, memory, or disconnected spreadsheets and provides a low-friction tool that can be installed from a web browser. The current iteration implements the operational core, while Chilean tax-compliant receipts and the production activation of Transbank remain outside its scope. The project integrates requirements analysis, data modeling, software development, security, and testing, making it relevant to the professional profile of an information technology graduate.')
para('Keywords: PWA, SaaS, neighborhood retail, inventory, point of sale, multi-tenant, artificial intelligence.',9,italic=True)

heading('1. Descripción y relevancia del Proyecto APT')
para('Localito funciona como una “caja inteligente de bolsillo”. Permite administrar varios negocios aislados entre sí y ofrece, según el rol, creación de locales y usuarios, catálogo de productos, stock y kardex, alertas, ventas con descuentos y pagos divididos, fiado, devoluciones, proveedores, órdenes de compra, caja por turno, auditoría e intercambio CSV. También incorpora una PWA instalable, cola local para operaciones sin conexión, lectura de códigos de barra y un flujo opcional de reconocimiento de envases mediante IA cuando existe una clave configurada.')
para('La relevancia laboral radica en que reproduce desafíos reales de la industria: levantamiento de necesidades, diseño de una solución escalable, separación segura de datos, integración frontend-backend-base de datos, reglas transaccionales, seguridad, pruebas y despliegue. Su beneficio esperado es mejorar la trazabilidad operativa de pequeños comercios y reducir errores de stock, caja y cuentas por cobrar. El ticket emitido por el MVP es interno y no reemplaza una boleta tributaria.')

heading('2. Relación con las competencias del perfil de egreso')
table(['Competencia','Aplicación en Localito','Evidencia verificable'],[
['Pruebas y certificación','Diseño de casos funcionales, automatización de reglas críticas y registro de pendientes manuales.','Matriz CP-01 a CP-46; pruebas automatizadas de sesiones, idempotencia, crédito, caja, compras e inventario.'],
['Gestión de proyectos informáticos','Priorización del MVP, alcance explícito, trabajo incremental, riesgos y decisiones de arquitectura.','Backlog, plan de trabajo, documentación de alcance y riesgos.'],
['Modelamiento de datos','Modelo relacional multi-negocio para usuarios, productos, movimientos, ventas, caja, compras y auditoría.','Esquema PostgreSQL y aislamiento por tenant derivado de la sesión.'],
['Desarrollo de soluciones de software','Construcción e integración de PWA React, API REST Node.js, PostgreSQL y servicios externos opcionales.','Código fuente, compilación, pruebas, endpoints y flujos demostrables.'],
['Implantación y mantenimiento','Configuración por variables de entorno, despliegue HTTPS y documentación de puesta en marcha.','README, scripts de ejecución y configuración para Vercel/PostgreSQL.']],[4.2,8.8,7])

heading('3. Relación con mis intereses profesionales')
para('Este proyecto se relaciona con mi interés por desarrollar soluciones digitales completas que resuelvan problemas cotidianos y tengan impacto visible en usuarios reales. Me permite profundizar en desarrollo full stack, aplicaciones móviles basadas en tecnologías web, diseño seguro de sistemas SaaS, automatización de pruebas e integración responsable de inteligencia artificial. También fortalece mi capacidad para traducir una necesidad de negocio en una arquitectura y un producto demostrable. La definición deberá complementarse con ejemplos personales del/de la estudiante durante la presentación oral, sin atribuir experiencias no acreditadas.')

heading('4. Factibilidad')
para('El proyecto es factible dentro de la asignatura porque el núcleo operacional ya se divide en módulos verificables y utiliza tecnologías disponibles: Node.js, React, PostgreSQL, navegador móvil y herramientas de desarrollo estándar. El alcance se controla mediante un MVP: la emisión tributaria ante el SII, múltiples sucursales complejas, facturación de la suscripción y Transbank productivo quedan fuera de esta iteración. El reconocimiento externo por IA es opcional; el sistema conserva un flujo controlado mediante código de barras o pista manual si no existe una clave de servicio.')
table(['Factor','Facilitador','Riesgo/obstáculo','Tratamiento'],[
['Tiempo','Arquitectura modular y núcleo ya implementado.','Ampliación excesiva del alcance.','Congelar MVP y priorizar defectos críticos y evidencia.'],
['Recursos','Herramientas abiertas, equipo personal y despliegue web.','Dependencia de PostgreSQL, HTTPS y servicios externos.','Modo demostración, configuración documentada y alternativas controladas.'],
['Técnica','Tipos compartidos, API y pruebas automatizadas.','Errores de integración, concurrencia u operación offline.','Pruebas de regresión, idempotencia, logs y matriz manual.'],
['Seguridad','Roles, sesiones, hashes y aislamiento por tenant.','Credenciales débiles o mal configuradas.','Variables privadas, rotación, validación y pruebas de autorización.'],
['Validación','Casos de prueba definidos.','Parte de la evidencia manual aún está pendiente.','Ejecutar casos priorizados, guardar capturas y registrar resultado real.']],[3,5,5,7])

heading('5. Objetivos')
para('Objetivo general',11,True)
para('Desarrollar y validar una PWA SaaS multi-negocio que apoye la gestión de ventas, inventario, caja, compras, clientes y fiados de pequeños comercios, incorporando reconocimiento de productos desde dispositivos móviles y controles de seguridad y trazabilidad.')
para('Objetivos específicos',11,True)
for x in ['Analizar y documentar las necesidades operativas y los límites del MVP.','Diseñar e implementar un modelo de datos escalable y aislado por negocio.','Integrar los módulos de autenticación, productos, stock, ventas, fiado, caja, compras y auditoría.','Incorporar lectura de códigos y reconocimiento visual opcional con confirmación humana.','Validar reglas críticas mediante pruebas automatizadas y ejecutar pruebas manuales funcionales y móviles.','Preparar evidencia técnica y una demostración reproducible, diferenciando funciones implementadas, simuladas y pendientes.']: bullet(x)

heading('6. Metodología de trabajo')
para('Se propone una metodología ágil e incremental basada en Scrum, adaptada a un proyecto académico individual. El backlog se ordena por valor y riesgo; cada iteración selecciona historias y criterios de aceptación, implementa un incremento, ejecuta revisión técnica y actualiza la evidencia. La calidad se aborda de manera transversal con control de versiones, revisión de alcance, validación de tipos, compilación, pruebas de API, pruebas manuales y registro de defectos. Para cada requisito se mantiene trazabilidad entre necesidad, módulo, caso de prueba y evidencia. Las integraciones externas se validan primero en modo de demostración o prueba y no se presentan como productivas sin credenciales ni confirmación real.')

heading('7. Plan de trabajo')
table(['Etapa','Actividades','Duración estimada','Recursos','Resultado'],[
['1. Definición','Problema, usuarios, alcance, objetivos, riesgos y backlog.','1 semana','Documentación y entrevistas/observación académica.','Definición aprobable y backlog priorizado.'],
['2. Diseño','Arquitectura, modelo de datos, seguridad, interfaz y criterios de aceptación.','1 semana','Diagramas, PostgreSQL, prototipos.','Diseño técnico trazable.'],
['3. Núcleo','Autenticación, roles, productos, inventario, ventas y fiado.','2 semanas','React, Node.js, repositorio.','Flujos centrales integrados.'],
['4. Operación','Caja, compras, devoluciones, CSV, PWA/offline y auditoría.','2 semanas','API, base de datos, navegador móvil.','MVP operacional.'],
['5. IA e integraciones','Código de barras, visión opcional y Webpay demo.','1 semana','Cámara, ZXing, servicio IA opcional.','Flujos demostrables y limitaciones documentadas.'],
['6. Validación','Typecheck, build, pruebas automatizadas y casos manuales priorizados.','1 semana','Matriz de pruebas y dispositivos.','Resultados y defectos registrados.'],
['7. Cierre','Correcciones, memoria, capturas, guion y ensayo.','1 semana','Documentación y entorno de demo.','Entrega y presentación reproducible.']],[2.7,7.6,2.7,4,4])

heading('8. Evidencias propuestas')
table(['Evidencia','Qué demuestra','Criterio de aceptación'],[
['Repositorio y estructura del código','Construcción e integración de componentes.','Módulos identificables, sin secretos versionados.'],
['Esquema y migración PostgreSQL','Modelo de datos y soporte multi-negocio.','Entidades y relaciones coherentes; RLS sin acceso público directo.'],
['Salida de typecheck, build y pruebas','Calidad técnica y regresión de reglas críticas.','Comandos terminan sin errores; resultados guardados con fecha.'],
['Matriz CP-01 a CP-46','Cobertura funcional y trazabilidad.','Cada caso indica estado real; no se aprueba sin ejecución o automatización.'],
['Capturas o video de flujos','Operación visible del MVP en escritorio/móvil.','Incluye rol, paso ejecutado y resultado observable.'],
['Demostración multi-tenant y permisos','Seguridad y separación de responsabilidades.','Vendedor recibe 403 en acción administrativa y no cambia tenant por cabecera.'],
['Documentación de alcance','Transparencia sobre funciones reales, simuladas y pendientes.','SII y Transbank productivo declarados fuera de alcance.']],[5,8,7])
para('A la fecha de esta definición, existen pruebas automatizadas aprobadas para reglas críticas; varios casos manuales permanecen “Pendiente evidencia”. Estos últimos deberán ejecutarse y respaldarse antes de afirmarlos como aprobados en la defensa.',9,italic=True)

heading('9. Indicadores de calidad disciplinarios')
table(['Indicador','Presencia en la propuesta','Estado/evidencia'],[
['1.1 Diseña pruebas de validación','Sí','Matriz funcional y criterios esperados definidos.'],
['1.2 Aplica pruebas de validación','Parcial','Automatizadas aprobadas; ejecución manual aún pendiente en varios casos.'],
['1.3 Desarrolla mejoras según resultados','Parcial','Flujo de corrección definido; requiere adjuntar bitácora de defectos/mejoras.'],
['2.1 Planifica proyectos informáticos','Sí','Alcance, etapas, recursos, riesgos y resultados.'],
['2.2 Controla proyectos informáticos','Parcial','Se propone seguimiento; debe conservarse evidencia periódica.'],
['3.1 Diseña modelos de datos','Sí','Modelo relacional multi-negocio documentado.'],
['3.2 Implementa modelos de datos','Sí','Esquema PostgreSQL implementado.'],
['4.1 Construye una solución de software','Sí','Núcleo operacional implementado.'],
['4.2 Integra componentes','Sí','PWA, API, datos y servicios opcionales integrados.'],
['4.3 Implanta una solución','Parcial','Despliegue preparado; operación productiva depende de configuración externa.']],[7,3,10])

heading('10. Autoevaluación según la rúbrica')
table(['IE','Nivel autoevaluado','Fundamento y mejora pendiente'],[
['1','Completamente logrado','La propuesta describe Localito, su impacto y relevancia profesional.'],['2','Completamente logrado','Relaciona cinco áreas del perfil con evidencias concretas.'],['3','Logrado','La relación profesional está explicitada; falta personalizar con experiencias acreditables del/de la estudiante.'],['4','Completamente logrado','Considera tiempo, recursos, riesgos, factores externos y mitigaciones.'],['5','Completamente logrado','Presenta un objetivo general y seis objetivos específicos coherentes.'],['6','Completamente logrado','Define trabajo ágil, incremental, trazable y orientado a calidad.'],['7','Completamente logrado','Incluye etapas, duración, recursos, riesgos y resultados.'],['8','Completamente logrado','Relaciona evidencias con logros y criterios de aceptación.'],['9','Completamente logrado','Redacción formal y referencias declaradas; requiere revisión final del estudiante.'],['10','Logrado','Se conserva la pauta institucional y se completa el informe; faltan datos personales.'],['11','Logrado','Diseño e implementación cubiertos; parte de validación, control e implantación mantiene evidencia pendiente.'],['12','Completamente logrado','Abstract, conclusions and reflection are written in English at an upper-intermediate level.']],[1.3,4.2,14.5])

heading('11. Conclusions',1)
para('Localito is a feasible and professionally relevant capstone project because it connects a real operational problem with an integrated software solution. Its current scope is controlled through a clear MVP and separates implemented features from simulated or future integrations. The project demonstrates competencies in data modeling, full-stack development, security, project planning, and software testing. Its strongest point is the operational breadth of the implemented core. The main remaining challenge is not to add more functions, but to complete the manual validation evidence, document defects and improvements, and rehearse a reproducible demonstration. These actions will strengthen the credibility of the final presentation and prevent unverified results from being reported as completed.')

heading('12. Reflection',1)
para('Working on Localito shows me that software quality depends on more than writing code. A useful solution needs a well-defined problem, realistic scope, secure data handling, testable requirements, and honest evidence. I have learned that external services such as payment platforms and AI models must be treated as dependencies with costs, credentials, privacy concerns, and fallback strategies. I also understand the importance of separating an academic demonstration from a production-ready service. My next priority is to execute the pending manual test cases, record the real outcomes, correct any defects, and explain the project in clear language for both technical and non-technical audiences.')

heading('Referencias',1)
para('Localito. (2026). Documento de Proyecto — Localito (versión 1.0) [Documento interno del proyecto].')
para('Localito. (2026). Matriz de pruebas funcionales — Localito [Documento interno del proyecto].')
para('Localito. (2026). README del repositorio Localito [Documentación técnica del proyecto].')

d.save(out)
