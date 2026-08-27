from docx import Document
import sys

doc = Document(sys.argv[1])
for i, p in enumerate(doc.paragraphs):
    print(f"P{i:03d} [{p.style.name}] {p.text!r}")
for ti, table in enumerate(doc.tables):
    print(f"\nTABLE {ti} rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows):
        vals=[]
        for ci, cell in enumerate(row.cells):
            vals.append(" / ".join(p.text.replace("\n", " | ") for p in cell.paragraphs))
        print(f"R{ri:03d}: " + " || ".join(vals))
for si, sec in enumerate(doc.sections):
    print(f"\nHEADER {si}")
    for p in sec.header.paragraphs: print(repr(p.text))
    for t in sec.header.tables:
        for r in t.rows: print(" || ".join(c.text.replace("\n", " | ") for c in r.cells))
    print(f"FOOTER {si}")
    for p in sec.footer.paragraphs: print(repr(p.text))
    for t in sec.footer.tables:
        for r in t.rows: print(" || ".join(c.text.replace("\n", " | ") for c in r.cells))
