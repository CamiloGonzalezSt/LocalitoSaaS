from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path
import os

NEW_REPORT = os.environ.get('LOCALITO_NEW_REPORT') == '1'
OUT = Path(os.environ.get('LOCALITO_REPORT_OUT', r"C:\Users\cajgo\Documents\Codex\LocalitoSaaS\entregables\1.4_APT122_FormativaFase1_COMPLETADA.docx"))
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "17365D"
LIGHT = "D9EAF7"
GRAY = "F2F2F2"
GREEN = "E2F0D9"
AMBER = "FFF2CC"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1.0 if NEW_REPORT else 0.75)
sec.bottom_margin = Inches(1.0 if NEW_REPORT else 0.7)
sec.left_margin = Inches(1.0 if NEW_REPORT else 0.85)
sec.right_margin = Inches(1.0 if NEW_REPORT else 0.85)
sec.header_distance = Inches(0.492 if NEW_REPORT else 0.3)
sec.footer_distance = Inches(0.492 if NEW_REPORT else 0.3)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'; normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
normal.paragraph_format.line_spacing = 1.5 if NEW_REPORT else 1.15
normal.paragraph_format.space_after = Pt(8 if NEW_REPORT else 6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
heading_tokens = [('Title',30,0,12),('Heading 1',16,18,10),('Heading 2',13,12,6),('Heading 3',12,8,4)] if NEW_REPORT else [('Title',24,0,12),('Heading 1',16,12,6),('Heading 2',13,10,4),('Heading 3',11,8,3)]
for name, size, before, after in heading_tokens:
    s=styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(BLUE)
    s._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial')
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=paragraph.add_run('Página ')
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE')
    r=OxmlElement('w:r'); t=OxmlElement('w:t'); t.text='1'; r.append(t); fld.append(r); paragraph._p.append(fld)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def style_table(table, widths=None, header=True):
    table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for ri,row in enumerate(table.rows):
        for ci,cell in enumerate(row.cells):
            set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: cell.width=Inches(widths[ci])
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.5 if NEW_REPORT else 1.0
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs: run.font.name='Arial'; run.font.size=Pt(11 if NEW_REPORT else 9.5)
            if header and ri==0:
                set_cell_shading(cell, 'F4F6F9' if NEW_REPORT else BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold=True
                        run.font.color.rgb=RGBColor.from_string(BLUE) if NEW_REPORT else RGBColor(255,255,255)
    if header and table.rows: set_repeat_table_header(table.rows[0])

def add_table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers))
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    for vals in rows:
        cells=t.add_row().cells
        for i,v in enumerate(vals): cells[i].text=str(v)
    style_table(t,widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def add_bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p

def add_number(text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p

def page_break(): doc.add_page_break()

# Header/footer
hp=sec.header.paragraphs[0]; hp.text='INFORME FORMATIVO · LOCALITO' if NEW_REPORT else 'LOCALITO | DEFINICIÓN PROYECTO APT'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.name='Arial'; r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
add_page_number(sec.footer.paragraphs[0])

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
r=p.add_run('LOCALITO'); r.font.name='Arial'; r.font.size=Pt(34); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PWA SaaS multi-negocio para comercios de barrio'); r.font.size=Pt(17); r.font.bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
r=p.add_run('Evaluación Formativa Fase 1\nDefinición de Proyecto APT'); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
doc.add_paragraph('\n')
meta=doc.add_table(rows=6, cols=2)
for i,(a,b) in enumerate([
    ('Asignatura','Capstone (PTY4614)'),('Integrantes','[NOMBRE INTEGRANTE 1] · [NOMBRE INTEGRANTE 2] · [NOMBRE INTEGRANTE 3, si corresponde]'),
    ('Docente','[NOMBRE DEL/DE LA DOCENTE]'),('Sección','[SECCIÓN]'),('Institución','Duoc UC'),('Fecha','[DD/MM/AAAA]')]):
    meta.rows[i].cells[0].text=a; meta.rows[i].cells[1].text=b
style_table(meta,[1.6,5.0],header=False)
set_repeat_table_header(meta.rows[0])
for row in meta.rows:
    set_cell_shading(row.cells[0],LIGHT)
    for rr in row.cells[0].paragraphs[0].runs: rr.font.bold=True; rr.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph('Documento preparado para revisión académica. Los campos entre corchetes deben completarse antes de entregar.')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
for r in p.runs: r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=RGBColor(90,90,90)

page_break()
doc.add_heading('Índice',0)
toc = [
('1. Resumen / Resumen ejecutivo',3),('2. Abstract',3),('3. Descripción y relevancia del Proyecto APT',4),
('4. Relación con el perfil de egreso',6),('5. Intereses profesionales',8),('6. Factibilidad',8),
('7. Desarrollo de ingeniería',11),('8. Calidad disciplinaria y estrategia de pruebas',17),
('9. Individual conclusions',19),('10. Reflection',20),('11. Bibliografía',21),('Anexo A. Trazabilidad y evidencia',22)
] if NEW_REPORT else [
('1. Resumen / Resumen ejecutivo',3),('2. Abstract',3),('3. Descripción y relevancia del Proyecto APT',4),
('4. Relación con el perfil de egreso',5),('5. Intereses profesionales',6),('6. Factibilidad',6),
('7. Desarrollo de ingeniería',8),('8. Calidad disciplinaria y estrategia de pruebas',12),
('9. Individual conclusions',13),('10. Reflection',13),('11. Bibliografía',14),('Anexo A. Trazabilidad y evidencia',15)]
for title,num in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.add_run(title); p.add_run('\t'+str(num))
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2))

page_break()
doc.add_heading('1. Resumen / Resumen ejecutivo',0)
doc.add_paragraph('Localito es una aplicación web progresiva (PWA) bajo un modelo SaaS multi-negocio, orientada a almacenes, minimarkets, botillerías, peluquerías y otros comercios de barrio. Centraliza punto de venta, inventario, caja, compras, proveedores, clientes, fiado y auditoría en una experiencia mobile-first. Su elemento diferenciador es el uso de la cámara del teléfono para leer códigos de barra o proponer coincidencias de productos mediante visión artificial cuando existe una clave de OpenAI configurada.')
doc.add_paragraph('El problema abordado es la fragmentación de la operación cotidiana: ventas registradas en cuadernos, stock conocido de memoria, deudas informales y escasa trazabilidad. Localito busca reducir errores, facilitar decisiones y entregar continuidad operacional desde un dispositivo disponible para el pequeño comerciante. El núcleo funcional está implementado, mientras que la emisión tributaria chilena y la validación manual completa permanecen fuera de esta fase.')
doc.add_heading('2. Abstract',0)
doc.add_heading('2.1 Resumen en español',1)
doc.add_paragraph('Localito propone una solución digital de bajo acceso para comercios de barrio mediante una PWA SaaS multi-negocio. El sistema integra ventas, inventario, clientes, crédito informal, caja, compras, proveedores y control de usuarios, con aislamiento de datos por negocio. Además, incorpora reconocimiento de productos por código de barras y un flujo opcional de visión artificial. La arquitectura utiliza React en la interfaz, Node.js con API REST en el backend y PostgreSQL como persistencia principal, con modo memoria para demostraciones controladas. La propuesta es relevante para el campo de la informática porque combina desarrollo full-stack, modelado de datos, seguridad, pruebas, despliegue y gestión ágil. El alcance actual prioriza el núcleo operacional y declara como pendientes la evidencia manual completa, la integración tributaria con SII y el reemplazo de Webpay simulado por una integración productiva.')
doc.add_heading('2.2 Abstract in English',1)
doc.add_paragraph('Localito is a low-barrier digital solution for neighborhood businesses delivered as a multi-tenant SaaS Progressive Web App. It integrates sales, inventory, customer credit, cash shifts, purchases, suppliers, and user management while isolating each business tenant. It also supports barcode recognition and an optional computer-vision workflow. The architecture combines a React client, a Node.js REST API, and PostgreSQL persistence, with an in-memory mode for controlled demonstrations. The project is relevant to the information technology field because it combines full-stack development, data modeling, security, testing, deployment, and agile project management. The current scope prioritizes the operational core and explicitly leaves full manual evidence collection, Chilean tax integration, and production Webpay integration for later stages.')

if not NEW_REPORT:
    page_break()
doc.add_heading('3. Descripción y relevancia del Proyecto APT',0)
doc.add_heading('3.1 Problema y oportunidad',1)
doc.add_paragraph('Los pequeños comercios suelen operar con herramientas dispersas o manuales. Esto dificulta conocer el stock real, controlar fiados, conciliar caja, reconstruir devoluciones y disponer de información confiable para decidir compras. La oportunidad consiste en entregar una herramienta simple, móvil y escalable que concentre esas funciones sin requerir infraestructura especializada en el local.')
if NEW_REPORT:
    page_break()
doc.add_heading('3.2 Solución propuesta',1)
add_table(['Área','Capacidad de Localito','Beneficio esperado'],[
('Ventas','Ticket, descuentos, notas, pagos simples/divididos e idempotencia','Atención más rápida y menor riesgo de cobros duplicados.'),
('Inventario','Productos, variantes, unidades, packs, stock mínimo, vencimientos y kardex','Mejor control de existencias y reposición.'),
('Clientes y fiado','Cupo, plazo, bloqueo, cuentas por cobrar, abonos y recordatorios','Trazabilidad del crédito informal.'),
('Caja y compras','Turnos, movimientos, cierre; proveedores, órdenes y recepción','Conciliación operativa y actualización de costos.'),
('Cámara e IA','ZXing para códigos y visión opcional comparada con el catálogo','Ingreso y consulta de productos con una interfaz natural.'),
('Administración','Roles system_admin, owner y seller; suspensión y auditoría','Gobierno multi-negocio y separación de responsabilidades.')],[1.2,3.0,2.4])
doc.add_heading('3.3 Relevancia para el campo laboral',1)
doc.add_paragraph('El proyecto reproduce problemas reales de la industria: autenticación segura, autorización por roles, aislamiento multi-tenant, consistencia de inventario, transacciones de venta, experiencia móvil, recuperación ante redes inestables y trazabilidad. Por ello permite demostrar competencias directamente transferibles a equipos de desarrollo de software, consultoría tecnológica, aseguramiento de calidad y gestión de productos digitales.')
doc.add_paragraph('El impacto es doble. Para el negocio, ordena procesos críticos con menor barrera de adopción. Para el equipo desarrollador, obliga a integrar interfaz, API, persistencia, seguridad y pruebas dentro de un alcance coherente. El ticket emitido es un comprobante interno y no sustituye una boleta tributaria; la integración con SII o un proveedor autorizado se mantiene fuera del MVP académico.')

if not NEW_REPORT:
    page_break()
doc.add_heading('4. Relación con las competencias del perfil de egreso',0)
add_table(['Competencia','Aplicación concreta en Localito','Evidencia/resultado'],[
('Realizar pruebas de certificación','Diseño de matriz funcional, pruebas automatizadas de reglas críticas y planificación de pruebas manuales en navegador y móvil.','Casos CP-01 a CP-46; siete casos automatizados figuran aprobados y el resto conserva estado pendiente de evidencia.'),
('Gestionar proyectos informáticos','Definición de alcance, backlog, sprints, riesgos, criterios de aceptación y prioridades de MVP.','Plan de nueve sprints (0–8), roadmap y mitigaciones documentadas.'),
('Construir modelos de datos','Modelo multi-tenant para negocios, usuarios, productos, stock, ventas, clientes, deudas, pagos, compras y auditoría.','Esquema PostgreSQL y modo memoria para demostración; tenant derivado desde la sesión.'),
('Desarrollar una solución de software','Implementación full-stack de una PWA, API REST, autenticación, roles, módulos operacionales, cola offline y despliegue preparado.','Código organizado en frontend/backend, scripts de compilación, pruebas y configuración de Vercel/PostgreSQL.')],[1.55,3.05,2.05])
if NEW_REPORT:
    page_break()
doc.add_heading('4.1 Indicadores de calidad disciplinarios',1)
for txt in [
'1.1 Diseño de pruebas: existe una matriz que relaciona módulos, pasos, resultado esperado y estado.',
'1.2 Aplicación de pruebas: hay pruebas automatizadas aprobadas para idempotencia, devoluciones, límite de crédito, caja, compras y recuperación de contraseña; las manuales no se presentan como aprobadas.',
'1.3 Mejora basada en resultados: la matriz funciona como línea base para registrar hallazgos, corregir defectos y repetir pruebas antes de la defensa.',
'2.1 Planificación: alcance, backlog, sprints, riesgos, criterios de aceptación y ambientes están definidos.',
'2.2 Control: los estados “Automatizada aprobada” y “Pendiente evidencia” permiten controlar avance sin sobrestimar resultados.',
'3.1 y 3.2 Datos: el modelo contempla relaciones, persistencia PostgreSQL, historial y aislamiento por negocio.',
'4.1 a 4.3 Software: la solución está construida e integrada; la implantación productiva completa depende de variables seguras, base persistente y validación final en dispositivos.'
]: add_bullet(txt)

page_break()
doc.add_heading('5. Relación con los intereses profesionales',0)
doc.add_paragraph('El proyecto se alinea con intereses profesionales en desarrollo full-stack, arquitectura de software, experiencia móvil, seguridad, datos y automatización de procesos. Localito permite trabajar sobre un producto completo y no sobre una función aislada: cada decisión de interfaz tiene efectos en la API, el modelo de datos, la trazabilidad y la experiencia del usuario.')
doc.add_paragraph('También ofrece una aproximación concreta al uso responsable de inteligencia artificial. La visión no reemplaza las reglas de negocio ni la confirmación humana: la imagen se reduce en el navegador, se procesa en el backend y la respuesta se compara con el catálogo del negocio. Cuando no hay clave configurada, se mantiene un flujo controlado por código o pista, evitando que una dependencia externa bloquee la demostración.')
doc.add_paragraph('Antes de entregar, cada integrante debe personalizar esta sección con una frase breve sobre su motivación, por ejemplo: “[NOMBRE]: mi interés profesional se centra en ______; mi aporte y aprendizaje dentro de Localito corresponde a ______”.')
doc.add_heading('6. Factibilidad',0)
doc.add_heading('6.1 Factibilidad técnica',1)
doc.add_paragraph('La base tecnológica es accesible para un proyecto académico: Node.js 20+, React, API REST, PostgreSQL 16 o Docker y herramientas habituales de control de versiones. La solución puede operar en modo memoria para una demostración rápida, aunque la persistencia real requiere PostgreSQL. La PWA evita depender de una aplicación nativa y permite probar en escritorio y teléfono.')
if NEW_REPORT:
    page_break()
doc.add_heading('6.2 Tiempo, materiales y factores externos',1)
add_table(['Dimensión','Condición','Respuesta de factibilidad'],[
('Tiempo','Asignatura organizada por sprints','Priorizar núcleo operacional y congelar alcance antes de pruebas finales.'),
('Materiales','PC, Node.js, navegador; PostgreSQL/Docker para persistencia','Herramientas disponibles y reproducibles mediante comandos documentados.'),
('Dispositivo','Teléfono para cámara y experiencia móvil','La demostración base funciona en navegador; cámara real requiere HTTPS o una red configurada.'),
('Servicios externos','OpenAI, correo, Vercel y base remota son configurables','Usar degradación controlada y no exponer claves en frontend o repositorio.'),
('Normativa','Boleta electrónica chilena fuera del MVP','Presentar solo ticket interno y declarar futura integración SII.'),
('Validación','Evidencia manual incompleta','Ejecutar matriz priorizada, capturar resultados y corregir antes de la defensa.')],[1.1,2.3,3.25])
page_break()
doc.add_heading('6.3 Riesgos y mitigaciones',1)
add_table(['Riesgo','Efecto','Mitigación'],[
('Alcance excesivo','Retraso o calidad irregular','Separar MVP, pendientes y roadmap; trabajar por prioridad.'),
('Conectividad inestable','Venta interrumpida o duplicada','Cola local, sincronización e idempotencia.'),
('Reconocimiento incorrecto','Producto equivocado','Código de barra primero, confianza visible y confirmación/corrección humana.'),
('Pérdida o cruce de datos','Impacto de seguridad','Sesiones seguras, autorización por rol y tenant derivado del token.'),
('Dependencia externa','Demo bloqueada','Flujos alternativos controlados y variables opcionales.'),
('Falta de evidencia','Resultado académico no demostrable','Matriz, capturas, registro de hallazgos y repetición de pruebas.')],[1.35,2.1,3.2])

page_break()
doc.add_heading('7. Desarrollo de ingeniería',0)
doc.add_heading('7.1 Objetivo general',1)
doc.add_paragraph('Diseñar e implementar una PWA SaaS multi-negocio que permita a comercios de barrio gestionar ventas, productos, stock, caja, compras, proveedores y clientes, incorporando fiado, operación offline controlada y reconocimiento asistido desde la cámara del teléfono.')
doc.add_heading('7.2 Objetivos específicos',1)
for txt in [
'Implementar autenticación, recuperación de contraseña, roles y aislamiento por negocio.',
'Centralizar el catálogo, stock, vencimientos, kardex, alertas y costos.',
'Registrar ventas con distintos medios de pago, devolución, anulación e idempotencia.',
'Administrar clientes, cupos, deudas, abonos y cobros compartibles.',
'Gestionar turnos de caja, compras, proveedores y auditoría.',
'Proveer una PWA instalable con comportamiento móvil y cola local ante pérdida de conexión.',
'Validar reglas críticas mediante pruebas automatizadas y completar evidencia manual priorizada.'
]: add_bullet(txt)
if NEW_REPORT:
    page_break()
doc.add_heading('7.3 Alcance implementado y límites',1)
add_table(['Estado','Contenido'],[
('Implementado','Núcleo operacional: plataforma/negocios, usuarios, sesiones, productos, stock, ventas, clientes y fiado, devoluciones, compras, caja, auditoría, CSV, offline/PWA, códigos de barra y visión opcional.'),
('Simulado/controlado','Webpay para demostración académica; reconocimiento por pista cuando no existe clave de visión.'),
('Pendiente de evidencia','Pruebas manuales de interfaz, móvil, PWA, permisos, cámara, red y flujos integrales indicadas en la matriz.'),
('Fuera de alcance','Boleta/factura electrónica SII, contabilidad completa, hardware fiscal y operación productiva sin una base persistente configurada.')],[1.45,5.2])
page_break()
doc.add_heading('7.4 Arquitectura',1)
add_table(['Capa','Responsabilidad','Tecnología'],[
('Presentación','Interfaz mobile-first, PWA, caché, cola local y captura de cámara.','React + TypeScript + Vite.'),
('Aplicación','API, autenticación, permisos, reglas de venta, stock, fiado, caja y compras.','Node.js + Express + TypeScript.'),
('Datos','Persistencia multi-negocio, auditoría y transacciones.','PostgreSQL; memoria solo para demo.'),
('Integraciones','Visión, correo, despliegue y pagos de demostración.','OpenAI opcional, Gmail/Resend, Vercel y Webpay simulado.')],[1.15,3.1,2.4])

if NEW_REPORT:
    page_break()
doc.add_heading('7.5 Modelo de datos y seguridad',0)
doc.add_paragraph('El negocio funciona como tenant. Las sesiones determinan el negocio y el rol; una cabecera proporcionada por el cliente no debe modificar ese contexto. Las entidades principales abarcan negocios, usuarios, sesiones, productos, movimientos de stock, ventas, ítems, pagos, clientes, cuentas por cobrar, abonos, proveedores, órdenes de compra, turnos de caja y registros de auditoría.')
add_table(['Control','Implementación o criterio'],[
('Contraseñas','Derivación mediante scrypt; no se almacenan en texto claro.'),('Sesiones','Tokens aleatorios almacenados como hash, con expiración y revocación.'),
('Autorización','Roles system_admin, owner y seller protegidos en interfaz y API.'),('Multi-tenant','El negocio se deriva desde la sesión; no se confía en x-tenant-id.'),
('Recuperación','Enlace de un solo uso, vencimiento de 30 minutos y revocación de sesiones anteriores.'),('Secretos','Claves solo en variables privadas; nunca con prefijo VITE_ ni en el repositorio.'),
('Auditoría','Registro de operaciones críticas para trazabilidad.')],[1.45,5.15])
page_break()
doc.add_heading('7.6 Flujos críticos',1)
for title,steps in [
('Venta',['Autenticar usuario y verificar turno/permiso.','Buscar o reconocer producto; validar disponibilidad.','Construir ticket, descuentos y pagos.','Confirmar con clave idempotente.','Registrar venta y movimientos; emitir ticket interno.']),
('Fiado',['Seleccionar cliente elegible.','Validar cupo, plazo y bloqueo.','Registrar venta y cuenta por cobrar.','Aceptar abonos y mantener saldo/historial.']),
('Compra',['Crear proveedor y orden.','Recibir cantidades.','Actualizar stock, kardex y costo promedio ponderado.']),
('Reconocimiento',['Intentar código de barra.','Si corresponde, enviar imagen reducida al backend.','Comparar propuesta con catálogo y mostrar confianza.','Solicitar confirmación o corrección.'])]:
    doc.add_heading(title,2)
    for s in steps: add_number(s)

page_break()
doc.add_heading('7.7 Gestión del proyecto',0)
doc.add_paragraph('Se adopta Scrum con entregas incrementales. El Product Backlog organiza historias y criterios de aceptación; cada Sprint Backlog selecciona trabajo verificable. La Definition of Done debe incluir código integrado, revisión de permisos, prueba pertinente, documentación mínima y ausencia de secretos expuestos.')
add_table(['Sprint','Objetivo principal'],[
('0','Preparación, arquitectura, backlog y repositorio.'),('1','Base técnica, autenticación y multi-tenant.'),('2','Productos e inventario.'),('3','Ventas y comprobante interno.'),('4','Clientes, fiado y abonos.'),('5','Cámara, códigos e IA.'),('6','Pagos de demostración y alertas.'),('7','Reportes, PWA y experiencia móvil.'),('8','Pruebas, correcciones, documentación y defensa.')],[1.0,5.6])
doc.add_heading('7.8 Puesta en marcha y verificación técnica',1)
doc.add_paragraph('Requisitos: Node.js 20 o superior, npm 10 o pnpm y PostgreSQL 16 o Docker Desktop para persistencia. Secuencia base:')
if NEW_REPORT:
    p=doc.add_paragraph(); p.style='Intense Quote'; p.paragraph_format.line_spacing=1.5; p.add_run('npm install · Copy-Item .env.example .env · npm run db:up · npm run dev:api · npm run dev:web')
else:
    for cmd in ['npm install','Copy-Item .env.example .env','npm run db:up','npm run dev:api','npm run dev:web']:
        p=doc.add_paragraph(); p.style='Intense Quote'; p.add_run(cmd)
doc.add_paragraph('Antes de presentar se deben ejecutar npm run typecheck, npm run build y npm run test -w apps/api. Esta entrega no infiere sus resultados: cualquier aprobación debe respaldarse con la salida real de esos comandos y con evidencia fechada.')

if not NEW_REPORT:
    page_break()
doc.add_heading('8. Calidad disciplinaria y estrategia de pruebas',0)
doc.add_heading('8.1 Enfoque',1)
doc.add_paragraph('La validación combina pruebas automatizadas para reglas de negocio críticas y pruebas manuales para experiencia, compatibilidad móvil, permisos visibles, cámara, PWA y comportamiento offline. Cada caso debe registrar ambiente, pasos, resultado observado, evidencia, responsable y fecha. Un caso pendiente no se considera aprobado hasta ejecutar y documentar su resultado.')
doc.add_heading('8.2 Estado verificable según la matriz',1)
add_table(['Estado','Casos','Interpretación'],[
('Automatizada aprobada','CP-35, CP-36, CP-37, CP-38, CP-39, CP-41 y CP-46.','La matriz registra aprobación automatizada de idempotencia, devoluciones, crédito, caja, compras y recuperación de contraseña.'),
('Pendiente evidencia','CP-01 a CP-34, CP-40 y CP-42 a CP-45.','Requieren ejecución/captura manual o evidencia adicional; no se declaran aprobados en este informe.')],[1.45,2.2,3.0])
doc.add_paragraph('Nota de consistencia: el documento fuente lista siete casos como “Automatizada aprobada”. La entrega conserva ese conteo y evita elevar a aprobado cualquier caso cuyo estado sea “Pendiente evidencia”.')
doc.add_heading('8.3 Plan de pruebas prioritario antes de la defensa',1)
add_table(['Prioridad','Casos','Criterio de salida'],[
('Alta','CP-01, CP-02, CP-07, CP-08, CP-24 a CP-26, CP-31 a CP-34','Flujos críticos y seguridad ejecutados sin defectos bloqueantes; capturas y respuestas API registradas.'),
('Alta','CP-22, CP-23, CP-44 y CP-45','Compatibilidad móvil/offline/visión validada en ambiente declarado; registrar limitaciones reales.'),
('Media','CP-03 a CP-06, CP-09 a CP-21, CP-27 a CP-30, CP-40, CP-42 y CP-43','Funciones administrativas y de soporte ejecutadas con evidencia.'),
('Regresión','CP-35 a CP-39, CP-41 y CP-46','Volver a ejecutar suite automatizada después de correcciones; conservar reporte de salida.')],[0.8,2.4,3.45])
doc.add_heading('8.4 Ciclo de mejora',1)
for txt in ['Ejecutar caso y capturar evidencia.','Registrar resultado observado y desviación.','Clasificar severidad y causa probable.','Corregir en una rama controlada.','Repetir el caso y la regresión asociada.','Actualizar la matriz solo con evidencia verificable.']: add_number(txt)

if NEW_REPORT:
    page_break()
doc.add_heading('9. Individual conclusions',0)
doc.add_heading('[Student 1 name]',1)
doc.add_paragraph('Localito demonstrates that a neighborhood-business problem can be addressed through an integrated software product rather than isolated screens. The project connects user experience, APIs, data consistency, security, and testing. Its main strength is the coherent operational core; its main challenge is to complete manual evidence and keep the scope controlled. My contribution and specific learning must be completed here before submission: [DESCRIBE CONTRIBUTION AND LEARNING].')
doc.add_heading('[Student 2 name]',1)
doc.add_paragraph('This project helped me understand how technical decisions affect real business processes. Multi-tenant isolation, idempotent sales, credit limits, and inventory movements are not only implementation details; they protect the reliability of the service. Before presenting the project, our team must validate the pending mobile and end-to-end scenarios and document the results honestly. My individual contribution must be completed here: [DESCRIBE CONTRIBUTION AND LEARNING].')
doc.add_heading('[Student 3 name, if applicable]',1)
doc.add_paragraph('Building Localito requires balancing innovation and feasibility. Computer vision adds value, but barcode scanning, confirmation, and fallback flows keep the product usable when an external service is unavailable. The project is feasible as an academic MVP because its boundaries are explicit and the critical features can be demonstrated independently. My individual contribution must be completed here: [DESCRIBE CONTRIBUTION AND LEARNING].')
if NEW_REPORT:
    page_break()
doc.add_heading('10. Reflection',0)
doc.add_paragraph('The first phase confirms that defining boundaries is as important as adding features. A credible project does not claim that every scenario is already validated: it separates implemented scope, automated evidence, manual work, simulated integrations, and future work. This distinction improves technical decision-making and academic transparency. The next step is to execute the prioritized test plan, collect reproducible evidence, correct defects, and rehearse a demonstration that clearly explains both the value of Localito and its current limitations.')

page_break()
doc.add_heading('11. Bibliografía y fuentes del proyecto',0)
sources=[
('Repositorio Localito.','README.md. Descripción funcional, requisitos, variables de entorno, puesta en marcha y verificación técnica. Consulta interna: 24-08-2026.'),
('Equipo Localito.','Documento-Proyecto-Localito.md. Definición, objetivos, alcance, arquitectura, datos, Scrum, riesgos y estado de desarrollo. Consulta interna: 24-08-2026.'),
('Equipo Localito.','Matriz-Pruebas-Localito.md. Casos funcionales, estados y evidencias recomendadas. Consulta interna: 24-08-2026.'),
('Duoc UC.','1.4_APT122_FormativaFase1.docx. Pauta de evaluación formativa de Definición Proyecto APT.')]
for author,desc in sources:
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Inches(-0.25); p.paragraph_format.left_indent=Inches(0.25)
    r=p.add_run(author+' '); r.font.bold=True; p.add_run(desc)
doc.add_heading('Declaración de alcance de las fuentes',1)
doc.add_paragraph('Las afirmaciones de implementación y pruebas provienen de los documentos internos indicados. Este informe no agrega resultados manuales inexistentes ni presenta como productivas las integraciones simuladas u opcionales.')

page_break()
doc.add_heading('Anexo A. Trazabilidad con la pauta y evidencia',0)
add_table(['Indicador de evaluación','Sección de respuesta','Evidencia prevista'],[
('1. Descripción y relevancia','Secciones 1, 2 y 3','Problema, solución, alcance, beneficios y relación con el campo laboral.'),
('2. Perfil de egreso','Sección 4','Matriz de competencias e indicadores disciplinarios.'),
('3. Intereses profesionales','Sección 5','Relación técnica y placeholders para personalización individual.'),
('4. Factibilidad','Sección 6','Tiempo, materiales, externos, riesgos y mitigaciones.'),
('5. Calidad disciplinaria','Secciones 7 y 8','Arquitectura, datos, seguridad, gestión, matriz y ciclo de mejora.')],[2.0,1.7,2.95])
doc.add_heading('A.1 Evidencias que deben adjuntarse después de ejecutarlas',1)
for txt in [
'Captura de inicio de sesión y dashboard.', 'Venta completa y ticket interno.', 'Stock antes/después de venta o devolución.',
'Fiado, abono y saldo actualizado.', 'Permisos de vendedor y rechazo 403 en API.', 'Caja: apertura, movimientos, cierre y diferencia.',
'PWA/celular, cámara o código de barras según ambiente.', 'Salida real de typecheck, build y pruebas automatizadas.',
'Registro de defectos corregidos y reejecución de casos afectados.'
]:
    p = add_bullet(txt)
    if NEW_REPORT:
        p.paragraph_format.space_after = Pt(0)
if not NEW_REPORT:
    doc.add_paragraph('Estado al emitir este documento: las evidencias manuales anteriores permanecen pendientes salvo que el equipo las ejecute, fecha y adjunte antes de entregar. No reemplazar este texto por “aprobado” sin respaldo.')

# Field refresh request and metadata
settings=doc.settings._element
upd=settings.find(qn('w:updateFields'))
if upd is None: upd=OxmlElement('w:updateFields'); settings.append(upd)
upd.set(qn('w:val'),'true')
doc.core_properties.title='Localito - Evaluación Formativa Fase 1'
doc.core_properties.subject='Definición Proyecto APT'
doc.core_properties.author='Equipo Localito'
doc.core_properties.keywords='Localito, PWA, SaaS, Capstone, APT'
doc.save(OUT)
print(OUT)
